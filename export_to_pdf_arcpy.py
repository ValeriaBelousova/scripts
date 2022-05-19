import arcpy
import requests
import os
import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

input_json = arcpy.GetParameterAsText(0)
#input_json = '[{"region_name": "Адыгея", "totalCount": 1, "totalSum": 1000000}, {"region_name": "Адыгея112", "totalCount": 1, "totalSum": 1000000}]'

title = 'Распределение по субъектам РФ'
cols_alias = {'region_name': u'Регион', 'totalCount': u'Количество', 'totalSum': u'Сумма'}

# json string to json
features_prop = json.loads(input_json)

# create docx
doc = docx.Document()

# create title
head = doc.add_heading(title)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

# create statistic table
table = doc.add_table(rows = len(features_prop), cols = len(cols_alias.keys()))
table.style = 'Table Grid'

# create heading cells
heading_cells = table.rows[0].cells

for i, heading in enumerate(list(cols_alias.values())):
    heading_cells[i].text = list(cols_alias.values())[i]
cells = table.add_row().cells

# fill table
for row in range(0, len(features_prop)):
    prop_count = len(features_prop[row])
    prop_list = list(features_prop[row].values())
    for col in range(0, prop_count):
        try:
            cell = table.cell(row + 1, col)
            cell.text = u'%s' % prop_list[col]
        except:
            pass

# save docx file
output_docx = os.path.join(arcpy.env.scratchFolder, '{}.docx'.format(title))
doc.save(output_docx)

# convert docx to pdf, remove docx file
convert(output_docx)
output_pdf = os.path.join(arcpy.env.scratchFolder, '{}.pdf'.format(title))

# Set the output parameter to be the output file of the server job
arcpy.SetParameterAsText(1, output_pdf)