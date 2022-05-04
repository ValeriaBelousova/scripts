import requests
import os
import json
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert

url = 'https://raw.githubusercontent.com/ValeriaBelousova/json/master/msk_distr.geojson?token=GHSAT0AAAAAABSLRLU42WHNKCRUOGMROQDEYTSHM3A'
response = requests.get(url)
json_response = response.json()

path = r'D:\__asp\__graph'
title = 'Распределение по субъектам РФ'
cols_alias = {'NAME': u'Название', 'NAME_EN': u'Название англ', 'ADMIN_LVL': u'Уровень'}


def getFeatures2Export(json_response):
    features_properties = []
    features = json_response["features"]
    print(len(features))
    for feature in features:
        properties = feature["properties"]
        features_properties.append(properties)
    return features_properties

def filterProperties(features_properties, fields_list):
    filter_feature_properties = []
    for properties in features_properties:
        filter_properties = {}
        for propertie in properties.keys():
            if propertie in fields_list:
                filter_properties[propertie] = properties[propertie]
        filter_feature_properties.append(filter_properties)          
    return filter_feature_properties

all_features_prop = getFeatures2Export(json_response)
filter_features_prop = filterProperties(all_features_prop, cols_alias.keys())

# create docx
doc = docx.Document()
head = doc.add_heading(title)
head.alignment = WD_ALIGN_PARAGRAPH.CENTER

table = doc.add_table(rows = len(filter_features_prop), cols = len(cols_alias.keys()))
table.style = 'Table Grid'
for row in range(1):
    for col in range(0, len(list(cols_alias.values()))):
        cell = table.cell(row, col)
        cell.text = list(cols_alias.values())[col]

for row in range(1, len(filter_features_prop)):
    prop_count = len(filter_features_prop[row])
    prop_list = list(filter_features_prop[row].values())
    for col in range(0, prop_count):
        try:
            cell = table.cell(row, col)
            cell.text = u'%s' % prop_list[col]
        except:
            pass
out_file = os.path.join(path, '%s.docx' % (title))
doc.save(out_file)
convert(out_file)